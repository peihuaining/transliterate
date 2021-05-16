#!/usr/bin/env python
# !-*-coding:utf-8 -*-
# !@Author:peihuaining

import re
import transliterate
from transliterate.discover import autodiscover
from transliterate.conf import set_setting
from transliterate import get_available_language_codes, translit
from hunspell import Hunspell
from docx import Document
from docx.shared import Inches

document = Document()
txt=["Yagaad Amerikchuud Mongolchuudad sanaa tavih bilee? Bidnii hoorond urag sadan, uuh tuuhiin yamar holboo baigaa bilee? Bodood bodood bodiin shiir 4 l bolood baidag??? Aa neeree bidend Arschilal geech saihan yum ogluu? Tegeed yadag bilee? Suuliin 30 jiliin Ih Hyamral nuurlej hamag baylag gadagshaa zoogdoj Mongolchuud toriin ajlaa tedenguigeer hiij chadahaa bolidog biluu??"
,"Haldwar awsan hun darhlaa togtdog gej yu yariad bgan beeee, vaccin bol haldwar awahgui gesen ug bish shu de, vaccin bol hunii bied todorhoi hemjeenii dsrhlaa togtoodog darhlaa ni herew ter hun haldwsr awlaa gehed vaccin hiilgeegui huniig bodohod hungun helberiin uvchulnu bas uheh ayulgui gesen ug, tehees vaccin hiilgesen bol haldwar awahgui gej buruu oilgoltoo hayarai,"
,"Ene deer uhwel taarna gej bichsen humuus yg yu hiij buteechiheed Zuckerbergiin achaar ardchillaar damjuulaad facebook tweetriin tsaanaas tom dugaraad baigaan uruwdeltei. "
,"yoooy ene hvmvvst gar hvrch boldoggvi burhad uu. hvmvvsiin uur buhimdal haana hvrch bgaag harj bgaa bizdee"
,"Ene chini nogoo tergentsereer mordongoos garsan nohor biluu"
,"YNZ BURIIN TAILBAR CH HEREGGUI.ERDENE ZUGEER L BOLICHIH. SER SER …….BAT UUL GEDEG ,ELBEG GEDEG ALUURCHIDTAIGAA ZUGEER L YV.ZA YU.ZORIG TA NARIIN GART UREGDSEN TER HAR UDRUUS MONGOLD ARDCHCHILAL BHGUISHDEE.ENE MINII L BODOL. GEHDEE UNENDEE BARAG.UDAHGUI MONGOL MYNMARAAS DOR TIIM ULS BOLNO.SONGOLTGUI SOGUULI GEJ ENIIG L HELDGIINSHDEE .YMAR NEGEN ULSIIN UNDSEN HUULI HEN NEGEN AVILGACHDIIN HUSLEER DURTAI TCAGTAA SHINECHILEGDDEG BVAL……"
,"ardchilsan nam gudamjind ulun deeremchin uussen nam terungeeree l duusaj baina"
,"Tanai namiin alban tushaal sandal shireenii asuudald MAN yamar hamaatai yum be. Nam dotroo uchraa olj chadahgui baij eronhiilogch boloh geed ulaaraad bh yum daa. Nam damjsan banzaluud ard tumniig turhirahaa bolio."
,"Ene altanhuyag gej neg bantan hutgaj har tagjgar pisda bainaa. Erh barij baihdaa oligtoihon baisan bol muu har lalarm min chi arai ingehgui bh bsanshuu llraa"
,"Dotroo asuudlaa shiydej chadahgui nam ulsiin asuudal shiydeh gej ulairhaa bolimoor yum"
,"gulugnuud chin tgd yah gj bgn sanaa ni zovood bgnmuu natogoo oruulaad ir 2 hurshuuruu nuhuulchie muu pizdaaa nar manaih Afghan irak ntr shig esvel Vietnam Myanmar ntr chin bish shuu davraad bgarai zolbin gulugnuud uhlee husej baival hureed ir muu umhii yankuud XAXAXAXAXA"
,"Hyatadin communist namin bodlogo chigleliig helbereltgui dagaj murdunu gedgee man in amarbayasgalan gej jduchin helsen yum bilee shd. Hurelee huurandaagiin uildel ch uuniig haruulj bgaa. Ter mahni uildveruudiig hjaad ugsun 700 horlon suitgegch troll oruulj irsen tahaltai hjaa narig hul hiriotoi bhad nuutsaar oruulj irsen.man tar."
,"hogiin amerikuud mongoliig bas samarch ongot hubisgalaa hiih gej bnaa.bas heden say dollar hayh biz.ukiraind yg ingej bij 5 terbumiig hayaad l samarch ogson doo.zvger l hogjiji baisan uls odoo dampuuraad duuslaa daa.. amerikiin arabiin hubisgal gej tunis livi siri ulsiig samarsan daa.bas gvrj kirgez armen ulsiig samarch hayaad bga.kirgezd bol heden jil bolood neg hubisgal hiinee.uls oron dampuuraad duuslaa da."
,"Za odoo amerikuud manaihiig daraagiin Irak, Ukrain bolgohoor hamgaalalh nereer idesh bolgoh ehlelee tavij baigaa yum bna daa! Kh. Battulga yerunhiilugch mine ee, huviinhaa erh ashgiin tuluu uls ornoo zoliosloj bui ene buzar, zavaan uildlee zogsoohgui yum uu?!!!"
,"Yagaad Amerikchuud Mongolchuudad sanaa tavih bilee? Bidnii hoorond urag sadan, uuh tuuhiin yamar holboo baigaa bilee? Bodood bodood bodiin shiir 4 l bolood baidag??? Aa neeree bidend Arschilal geech saihan yum ogluu? Tegeed yadag bilee? Suuliin 30 jiliin Ih Hyamral nuurlej hamag baylag gadagshaa zoogdoj Mongolchuud toriin ajlaa tedenguigeer hiij chadahaa bolidog biluu??"
,"Man hyatadin ko namin daguul bolson n huukiin hiisen uldel amarbayasgalangiin helsen ug notolj bna shdee. Huukiig bhad mahni uildveruud hujaad ochson mun tahaltai hjaa oruulj irj tahal taraasan darhni zamig tam bolgoson hyatadaas mash ih zeel avsan. Idej uusnig m yarihaa boly. Oyuerdene garch ireed demb huleen avaagui hyatad vaktsinig oruulj irj ard tumnee turshiltand oruulsan odoo bur huchindej hiigeed uvsad 2 tsagdaa uhchlee gej bna. Ene uheed bgaa humuusiin hed n uunees bolsong hen medlee man hkn 2 meduuleh ch ugui biz. Hul horio n sulranguut talbai orno. Mongol mongolooroo uldehiig husch bgaa mongol hun bur talbai oroh heregtei."
]
def seg_tail_translit_mn_split(str1,Ischeck,sep=r',|\.|/|;|\'|`|\[|\]|<|>|\?|:|"|\{|\}|\~|!|@|#|\$|%|\^|&|\(|\)|-|=|\_|\+|，|。|、|；|‘|’|【|】|·|！|…|（|）'): # 分隔符可为多样的正则表达式
    # 保留分割符号，置于句尾，比如标点符号
    #Ischeck (True or False) 是否进行字典纠错
    wlist = re.split(sep,str1)
    #纠错
    h = Hunspell('mongolian', hunspell_data_dir='./dict/mn/')
    res=[]
    for tempstr in wlist:
        temps=[]
        tempstr = translit(tempstr.lower(), 'mn').strip()
        for word in tempstr.split():
            if Ischeck and not h.spell(word) and len(h.suggest(word))!=0:
                temps.append(h.suggest(word)[0])
            else :
                #print(h.suggest(word))
                temps.append(word)
        res.append(' '.join(temps))
    wlist = res
    seg_word = re.findall(sep,str1)
    seg_word.extend(" ") # 末尾插入一个空字符串，以保持长度和切割成分相同
    wlist = [ x+y for x,y in zip(wlist,seg_word) ] # 顺序可根据需求调换
    return wlist

#添加标题，并设置级别，范围：0 至 9，默认为1
document.add_heading('蒙古文测试——带纠错', 0)

for tx in txt:
    document.add_paragraph(tx)
    print(tx)
    res = ' '.join(seg_tail_translit_mn_split(tx,True))
    print(res)
    document.add_paragraph(res)

document.save('demo.docx')




